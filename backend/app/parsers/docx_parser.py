"""
Word document parser using python-docx.
Extracts text with section headings and estimated page numbers.
"""

from typing import List, Tuple
from io import BytesIO

from docx import Document


class DocxParser:
    """Parse Word documents and extract text with section information."""

    # Approximate characters per page for page estimation
    CHARS_PER_PAGE_ESTIMATE = 3000

    def parse(self, file_content: bytes, filename: str) -> List[Tuple[int, str, str]]:
        """
        Parse DOCX and extract text with section headings.

        Args:
            file_content: Raw DOCX file bytes
            filename: Original filename (for logging)

        Returns:
            List of tuples: (page_number, section_heading, text_content)
            Note: page_number is estimated based on character count
        """
        doc = Document(BytesIO(file_content))
        results = []

        current_heading = ""
        current_text = ""
        char_count = 0
        page_num = 1

        for para in doc.paragraphs:
            # Detect headings (Heading 1, Heading 2, etc.)
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # Save previous section if it has content
                if current_text.strip():
                    results.append((page_num, current_heading, current_text.strip()))

                current_heading = para.text
                current_text = ""
            else:
                current_text += para.text + "\n"
                char_count += len(para.text)

                # Estimate page breaks
                if char_count >= self.CHARS_PER_PAGE_ESTIMATE:
                    page_num += 1
                    char_count = 0

        # Don't forget the last section
        if current_text.strip():
            results.append((page_num, current_heading, current_text.strip()))

        # If no content was extracted (no paragraphs), return empty
        if not results:
            # Try to get any text from the document
            all_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if all_text:
                results.append((1, "", all_text))

        return results

    def get_page_count(self, file_content: bytes) -> int:
        """
        Estimate page count based on character count.

        Note: DOCX files don't have a built-in page count property,
        so this is an approximation.
        """
        doc = Document(BytesIO(file_content))
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        return max(1, (total_chars + self.CHARS_PER_PAGE_ESTIMATE - 1) // self.CHARS_PER_PAGE_ESTIMATE)
